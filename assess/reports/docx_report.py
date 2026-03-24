"""DOCX report generation using python-docx."""

from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..config import CHECKS, RECOMMENDATIONS
from ..engine import readiness_level
from . import prepare_report_data


def _level_color(level):
    return {
        "Ready": RGBColor(0x10, 0xB9, 0x81),
        "Partially Ready": RGBColor(0xF5, 0x9E, 0x0B),
        "Needs Work": RGBColor(0xF9, 0x73, 0x16),
        "Not Ready": RGBColor(0xEF, 0x44, 0x44),
    }.get(level, RGBColor(0, 0, 0))


def _score_color(score):
    if score >= 80:
        return RGBColor(0x10, 0xB9, 0x81)
    if score >= 60:
        return RGBColor(0xF5, 0x9E, 0x0B)
    if score >= 40:
        return RGBColor(0xF9, 0x73, 0x16)
    return RGBColor(0xEF, 0x44, 0x44)


def _set_cell_text(cell, text, bold=False, color=None, size=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=Pt(9))

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if isinstance(val, tuple):
                text, color = val
                _set_cell_text(cell, text, color=color, size=Pt(9))
            else:
                _set_cell_text(cell, val, size=Pt(9))

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def generate_docx(all_results, title="AI Bug Automation Readiness Report", org=None):
    if not all_results:
        doc = Document()
        doc.add_heading("No repositories found to assess.", level=1)
        return doc

    d = prepare_report_data(all_results, org=org)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    org_prefix = d["org_prefix"]

    # --- Build document ---
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"{len(d['sorted_results'])} repositories assessed -- {now}")

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"{d['ready_count']} of {len(d['sorted_results'])}").bold = True
    p.add_run(f" repositories are partially ready or above for AI-assisted bug fixing. "
              f"The ecosystem averages ")
    score_run = p.add_run(f"{d['avg']:.0f}/100")
    score_run.bold = True
    score_run.font.color.rgb = _score_color(d["avg"])
    p.add_run(f". The biggest gap is \"{d['biggest_gap'][1]}\" (avg {d['biggest_gap'][2]:.0f}/100).")

    # How scoring works
    doc.add_heading("How Scoring Works", level=1)
    doc.add_paragraph(
        f"Overall score = weighted average of {len(CHECKS)} checks, each scored 0-100."
    )
    doc.add_paragraph(
        f"Phases: Understand ({d['cat_weights'].get('Understand', 0)}%), "
        f"Navigate ({d['cat_weights'].get('Navigate', 0)}%), "
        f"Verify ({d['cat_weights'].get('Verify', 0)}%), "
        f"Submit ({d['cat_weights'].get('Submit', 0)}%)."
    )
    doc.add_paragraph(
        "Verify phase gate: If the average Verify score is below 50, "
        "the overall score receives a smooth penalty multiplier that scales "
        "linearly from x0.4 (verify avg = 0) to x1.0 (verify avg = 50)."
    )
    _add_table(doc,
               ["Level", "Score", "Meaning"],
               [
                   [("Ready", _level_color("Ready")), "80+",
                    "Strong test infrastructure, good context, CI validates fixes"],
                   [("Partially Ready", _level_color("Partially Ready")), "60-79",
                    "Workable but has gaps that may cause AI agent failures"],
                   [("Needs Work", _level_color("Needs Work")), "40-59",
                    "Missing key capabilities; fix gaps before AI bug bash"],
                   [("Not Ready", _level_color("Not Ready")), "<40",
                    "Fundamental gaps in test infrastructure or code structure"],
               ],
               col_widths=[1.5, 0.8, 4.2])

    # Summary table
    doc.add_heading("Summary", level=1)
    _add_table(doc,
               ["Metric", "Value"],
               [
                   ["Average Score", f"{d['avg']:.0f}/100"],
                   [("Ready (80+)", _level_color("Ready")), str(d["tier_counts"]["Ready"])],
                   [("Partially Ready (60-79)", _level_color("Partially Ready")),
                    str(d["tier_counts"]["Partially Ready"])],
                   [("Needs Work (40-59)", _level_color("Needs Work")),
                    str(d["tier_counts"]["Needs Work"])],
                   [("Not Ready (<40)", _level_color("Not Ready")),
                    str(d["tier_counts"]["Not Ready"])],
               ],
               col_widths=[3.0, 1.5])

    # Phase scores
    doc.add_heading("Phase Scores (Ecosystem Average)", level=1)
    phase_rows = []
    for phase in ["Understand", "Navigate", "Verify", "Submit"]:
        pavg = d["cat_avgs"].get(phase, 0)
        phase_rows.append([phase, f"{d['cat_weights'].get(phase, 0)}%",
                           (f"{pavg:.0f}/100", _score_color(pavg))])
    _add_table(doc, ["Phase", "Weight", "Average"], phase_rows,
               col_widths=[2.0, 1.0, 1.5])

    # Bug bash shortlist
    doc.add_heading("Bug Bash Shortlist", level=1)

    for tier_label, tier_repos, show_gap in [
        ("Ready (80+)", d["ready_repos"], False),
        ("Partially Ready (60-79)", d["partially_ready"], True),
        ("Needs Work (40-59)", d["needs_work"], True),
        ("Not Ready (<40)", d["not_ready"], False),
    ]:
        doc.add_heading(f"{tier_label} -- {len(tier_repos)} repos", level=2)
        if not tier_repos:
            doc.add_paragraph("None")
            continue
        for r in tier_repos:
            name = f"{org_prefix}{r['repo']}" if org else r["repo"]
            text = f"{name} ({round(r['overall_score'])})"
            if show_gap:
                checks_sorted = sorted(
                    [(k, v) for k, v in r["checks"].items() if not v.get("excluded")],
                    key=lambda x: (100 - x[1]["score"]) * x[1]["weight"],
                    reverse=True,
                )
                top_gap = checks_sorted[0][1]["name"] if checks_sorted else ""
                text += f" -- top gap: {top_gap}"
            doc.add_paragraph(text, style="List Bullet")

    # Repository scores table
    doc.add_heading("Repository Scores", level=1)
    repo_rows = []
    for r in d["sorted_results"]:
        level, _ = readiness_level(r["overall_score"])
        langs = ", ".join(r["languages"][:3])
        gate = " (verify gate)" if r.get("verify_gate") else ""
        name = f"{org_prefix}{r['repo']}" if org else r["repo"]
        repo_rows.append([
            name + gate,
            (f"{round(r['overall_score'])}/100", _score_color(r["overall_score"])),
            (level, _level_color(level)),
            langs,
        ])
    _add_table(doc, ["Repository", "Score", "Level", "Languages"], repo_rows,
               col_widths=[2.5, 1.0, 1.5, 1.5])

    # Quick wins
    if d["quick_wins"]:
        doc.add_heading("Quick Wins (Highest Impact Actions)", level=1)
        qw_rows = []
        for name, repos_below, lift, rec, weight in d["quick_wins"][:7]:
            qw_rows.append([name, f"{repos_below} repos", f"{weight}%", rec])
        _add_table(doc, ["Action", "Repos Below 40", "Weight", "How to Fix"], qw_rows,
                   col_widths=[1.8, 1.0, 0.7, 3.0])

    # All checks ranked
    doc.add_heading("All Checks Ranked by Average Score", level=1)
    check_rows = []
    for cid, name, avg_score, weight in d["worst_checks"]:
        check_rows.append([name, (f"{avg_score:.0f}/100", _score_color(avg_score)),
                           f"{weight}%"])
    _add_table(doc, ["Check", "Avg Score", "Weight"], check_rows,
               col_widths=[3.0, 1.5, 1.0])

    # Per-repository details
    doc.add_heading("Per-Repository Details", level=1)
    check_ids = list(CHECKS.keys())
    for r in d["sorted_results"]:
        level, _ = readiness_level(r["overall_score"])
        profile_suffix = ""
        if r.get("profile", {}).get("name", "default") != "default":
            p = r["profile"]
            profile_suffix = f" [profile: {p['name']}]"
        doc.add_heading(
            f"{r['repo']} -- {round(r['overall_score'])}/100 ({level}){profile_suffix}", level=2
        )
        detail_rows = []
        for cid in check_ids:
            c = r["checks"][cid]
            if c.get("excluded"):
                grey = RGBColor(0x9C, 0xA3, 0xAF)
                detail_rows.append([
                    c["name"],
                    c["category"],
                    f"{c['weight']}%",
                    ("N/A", grey),
                    "Excluded by profile",
                ])
            else:
                ev_text = "; ".join(c["evidence"])
                if c.get("recommendation"):
                    ev_text += f" Rec: {c['recommendation']}"
                detail_rows.append([
                    c["name"],
                    c["category"],
                    f"{c['weight']}%",
                    (f"{c['score']:.0f}", _score_color(c["score"])),
                    ev_text,
                ])
        _add_table(doc, ["Check", "Phase", "Weight", "Score", "Evidence"], detail_rows,
                   col_widths=[1.5, 0.8, 0.6, 0.6, 3.0])

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"AI Bug Automation Readiness Report -- Generated {now}")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc
